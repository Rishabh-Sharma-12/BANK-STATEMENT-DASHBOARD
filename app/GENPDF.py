from docx import Document
from io import BytesIO

def generate_docx(question, response):
    doc = Document()
    doc.add_heading('AI Financial Analysis Report', 0)
    doc.add_heading('Question:', level=1)
    doc.add_paragraph(question)
    doc.add_heading('Response:', level=1)
    doc.add_paragraph(response)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer